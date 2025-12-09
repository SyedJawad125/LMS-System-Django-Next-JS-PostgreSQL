# ============================================
# OPTIMIZED DOCUMENT READER
# File: apps/rag_system/services/pdf_reader.py
# ============================================

import PyPDF2
import os
from typing import Optional, Dict

try:
    from docx import Document
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False
    print("⚠️ python-docx not installed. DOCX reading will not be available.")


class DocumentReader:
    """Enhanced document reader with better error handling"""
    
    SUPPORTED_FORMATS = ['pdf', 'docx', 'txt', 'csv', 'md']
    
    @staticmethod
    def read_pdf(file_path: str) -> str:
        """
        Extract text from PDF with improved error handling
        
        Args:
            file_path: Path to PDF file
            
        Returns:
            Extracted text or empty string on error
        """
        try:
            if not os.path.exists(file_path):
                print(f"❌ PDF file not found: {file_path}")
                return ""
            
            text_parts = []
            
            with open(file_path, 'rb') as file:
                pdf_reader = PyPDF2.PdfReader(file)
                
                total_pages = len(pdf_reader.pages)
                print(f"📄 Reading PDF: {total_pages} pages")
                
                for page_num, page in enumerate(pdf_reader.pages, 1):
                    try:
                        page_text = page.extract_text()
                        if page_text:
                            # Clean up text
                            page_text = page_text.strip()
                            if page_text:
                                text_parts.append(page_text)
                    except Exception as e:
                        print(f"⚠️ Error reading page {page_num}: {e}")
                        continue
                
                full_text = "\n\n".join(text_parts)
                print(f"✅ Extracted {len(full_text)} characters from PDF")
                return full_text
                
        except Exception as e:
            print(f"❌ Error reading PDF {file_path}: {e}")
            return ""
    
    @staticmethod
    def read_docx(file_path: str) -> str:
        """
        Extract text from Word document
        
        Args:
            file_path: Path to DOCX file
            
        Returns:
            Extracted text or empty string on error
        """
        if not DOCX_AVAILABLE:
            print("❌ python-docx not installed. Install with: pip install python-docx")
            return ""
        
        try:
            if not os.path.exists(file_path):
                print(f"❌ DOCX file not found: {file_path}")
                return ""
            
            doc = Document(file_path)
            
            # Extract paragraphs
            paragraphs = [para.text for para in doc.paragraphs if para.text.strip()]
            
            # Extract tables if any
            table_texts = []
            for table in doc.tables:
                for row in table.rows:
                    row_text = " | ".join([cell.text.strip() for cell in row.cells if cell.text.strip()])
                    if row_text:
                        table_texts.append(row_text)
            
            # Combine all text
            all_text = "\n\n".join(paragraphs)
            if table_texts:
                all_text += "\n\nTABLES:\n" + "\n".join(table_texts)
            
            print(f"✅ Extracted {len(all_text)} characters from DOCX")
            return all_text.strip()
            
        except Exception as e:
            print(f"❌ Error reading DOCX {file_path}: {e}")
            return ""
    
    @staticmethod
    def read_txt(file_path: str) -> str:
        """
        Read plain text file with multiple encoding support
        
        Args:
            file_path: Path to text file
            
        Returns:
            File content or empty string on error
        """
        try:
            if not os.path.exists(file_path):
                print(f"❌ Text file not found: {file_path}")
                return ""
            
            # Try multiple encodings
            encodings = ['utf-8', 'latin-1', 'cp1252', 'iso-8859-1']
            
            for encoding in encodings:
                try:
                    with open(file_path, 'r', encoding=encoding) as file:
                        text = file.read().strip()
                        print(f"✅ Read {len(text)} characters from TXT (encoding: {encoding})")
                        return text
                except UnicodeDecodeError:
                    continue
            
            # If all encodings fail, try binary mode
            print("⚠️ All text encodings failed, trying binary mode")
            with open(file_path, 'rb') as file:
                content = file.read()
                # Try to decode as UTF-8, replacing errors
                text = content.decode('utf-8', errors='replace').strip()
                print(f"✅ Read {len(text)} characters from TXT (binary mode)")
                return text
                
        except Exception as e:
            print(f"❌ Error reading TXT {file_path}: {e}")
            return ""
    
    @staticmethod
    def read_csv(file_path: str) -> str:
        """
        Read CSV file and convert to text format
        
        Args:
            file_path: Path to CSV file
            
        Returns:
            CSV content as formatted text
        """
        try:
            if not os.path.exists(file_path):
                print(f"❌ CSV file not found: {file_path}")
                return ""
            
            with open(file_path, 'r', encoding='utf-8') as file:
                lines = file.readlines()
                
                # Format CSV as readable text
                text_lines = []
                for i, line in enumerate(lines):
                    line = line.strip()
                    if line:
                        if i == 0:
                            text_lines.append(f"HEADERS: {line}")
                        else:
                            text_lines.append(f"Row {i}: {line}")
                
                full_text = "\n".join(text_lines)
                print(f"✅ Read {len(lines)} rows from CSV")
                return full_text
                
        except Exception as e:
            print(f"❌ Error reading CSV {file_path}: {e}")
            return ""
    
    @staticmethod
    def read_markdown(file_path: str) -> str:
        """
        Read Markdown file
        
        Args:
            file_path: Path to MD file
            
        Returns:
            Markdown content
        """
        # Markdown is just text, use text reader
        return DocumentReader.read_txt(file_path)
    
    @classmethod
    def read_document(cls, file_path: str, doc_type: str = None) -> str:
        """
        Read document based on type with auto-detection
        
        Args:
            file_path: Path to document
            doc_type: Document type (pdf, docx, txt, csv, md) or None for auto-detect
            
        Returns:
            Extracted text or empty string on error
        """
        try:
            # Auto-detect if not specified
            if not doc_type:
                _, ext = os.path.splitext(file_path)
                doc_type = ext.lstrip('.').lower()
            
            doc_type = doc_type.lower()
            
            print(f"\n📖 Reading document: {os.path.basename(file_path)}")
            print(f"   Type: {doc_type}")
            
            # Route to appropriate reader
            if doc_type == 'pdf':
                return cls.read_pdf(file_path)
            elif doc_type == 'docx' or doc_type == 'doc':
                return cls.read_docx(file_path)
            elif doc_type == 'txt':
                return cls.read_txt(file_path)
            elif doc_type == 'csv':
                return cls.read_csv(file_path)
            elif doc_type == 'md' or doc_type == 'markdown':
                return cls.read_markdown(file_path)
            else:
                print(f"❌ Unsupported document type: {doc_type}")
                print(f"   Supported formats: {', '.join(cls.SUPPORTED_FORMATS)}")
                return ""
                
        except Exception as e:
            print(f"❌ Error in read_document: {e}")
            return ""
    
    @classmethod
    def get_document_info(cls, file_path: str) -> Dict:
        """
        Get information about a document without reading full content
        
        Args:
            file_path: Path to document
            
        Returns:
            Dictionary with document metadata
        """
        try:
            if not os.path.exists(file_path):
                return {"error": "File not found"}
            
            file_size = os.path.getsize(file_path)
            _, ext = os.path.splitext(file_path)
            doc_type = ext.lstrip('.').lower()
            
            info = {
                "filename": os.path.basename(file_path),
                "file_type": doc_type,
                "file_size_bytes": file_size,
                "file_size_mb": round(file_size / (1024 * 1024), 2),
                "exists": True,
                "supported": doc_type in cls.SUPPORTED_FORMATS
            }
            
            # Get page count for PDFs
            if doc_type == 'pdf':
                try:
                    with open(file_path, 'rb') as file:
                        pdf_reader = PyPDF2.PdfReader(file)
                        info["page_count"] = len(pdf_reader.pages)
                except:
                    info["page_count"] = "Unknown"
            
            # Get paragraph count for DOCX
            elif doc_type == 'docx' and DOCX_AVAILABLE:
                try:
                    doc = Document(file_path)
                    info["paragraph_count"] = len(doc.paragraphs)
                    info["table_count"] = len(doc.tables)
                except:
                    pass
            
            # Get line count for text files
            elif doc_type in ['txt', 'csv', 'md']:
                try:
                    with open(file_path, 'r', encoding='utf-8') as file:
                        info["line_count"] = len(file.readlines())
                except:
                    pass
            
            return info
            
        except Exception as e:
            return {"error": str(e)}
    
    @classmethod
    def validate_document(cls, file_path: str) -> tuple[bool, str]:
        """
        Validate if document can be read
        
        Args:
            file_path: Path to document
            
        Returns:
            Tuple of (is_valid, error_message)
        """
        try:
            # Check file exists
            if not os.path.exists(file_path):
                return False, "File not found"
            
            # Check file size (limit to 50MB)
            file_size = os.path.getsize(file_path)
            max_size = 50 * 1024 * 1024  # 50MB
            if file_size > max_size:
                return False, f"File too large: {file_size / (1024*1024):.2f}MB (max 50MB)"
            
            # Check file type
            _, ext = os.path.splitext(file_path)
            doc_type = ext.lstrip('.').lower()
            
            if doc_type not in cls.SUPPORTED_FORMATS:
                return False, f"Unsupported format: {doc_type}"
            
            # Check DOCX availability
            if doc_type == 'docx' and not DOCX_AVAILABLE:
                return False, "python-docx not installed"
            
            return True, "Valid"
            
        except Exception as e:
            return False, str(e)


# ============================================
# TESTING FUNCTIONS
# ============================================

def test_document_reader():
    """Test the document reader with sample files"""
    print("\n" + "="*60)
    print("🧪 Testing Document Reader")
    print("="*60)
    
    reader = DocumentReader()
    
    # Test file info
    print("\n1. Testing document info...")
    test_file = "/path/to/test.pdf"  # Update with actual path
    info = reader.get_document_info(test_file)
    print(f"   File Info: {info}")
    
    # Test validation
    print("\n2. Testing document validation...")
    is_valid, message = reader.validate_document(test_file)
    print(f"   Valid: {is_valid}, Message: {message}")
    
    # Test reading
    print("\n3. Testing document reading...")
    text = reader.read_document(test_file)
    print(f"   Extracted {len(text)} characters")
    print(f"   Preview: {text[:200]}...")
    
    print("\n" + "="*60)
    print("✅ Testing complete!")
    print("="*60)


if __name__ == "__main__":
    test_document_reader()